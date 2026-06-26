import os
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


class ReportGenerator:

    def __init__(self):
        self.doc = Document()

    def add_title(self, text):
        self.doc.add_heading(text, level=0)

    def add_heading(self, text, level=1):
        self.doc.add_heading(text, level=level)

    def add_paragraph(self, text):
        self.doc.add_paragraph(text)

    def add_table(self, headers, rows):
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Shading Accent 1'
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for row_data in rows:
            row = table.add_row()
            for i, val in enumerate(row_data):
                row.cells[i].text = str(val)

    def add_image(self, img_path, width=5.5):
        self.doc.add_picture(img_path, width=Inches(width))

    def plot_histogram(self, data, title, xlabel, ylabel, save_path):
        plt.figure(figsize=(8, 5))
        plt.hist(data, bins=20, edgecolor='black', alpha=0.7)
        plt.title(title)
        plt.xlabel(xlabel)
        plt.ylabel(ylabel)
        plt.grid(True, alpha=0.3)
        plt.tight_layout()
        plt.savefig(save_path, dpi=150)
        plt.close()
        return save_path

    def plot_cumulative(self, data, title, xlabel, ylabel, save_path):
        plt.figure(figsize=(8, 5))
        sorted_data = np.sort(data)
        cumsum = np.arange(1, len(sorted_data) + 1) / len(sorted_data)
        plt.plot(sorted_data, cumsum, 'b-', linewidth=2)
        plt.title(title)
        plt.xlabel(xlabel)
        plt.ylabel(ylabel)
        plt.grid(True, alpha=0.3)
        plt.tight_layout()
        plt.savefig(save_path, dpi=150)
        plt.close()
        return save_path

    def save(self, path):
        self.doc.save(path)
